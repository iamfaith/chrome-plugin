import pdfquery
import pandas as pd

pdf_file = "/home/faith/faith/Downloads/chizong1.pdf"
# pdf = pdfquery.PDFQuery(pdf_file)
# pdf.load()

# pdf.tree.write('patient1xml.xml', pretty_print = True)
# print(pdf)

# questions = [] # create an empty list to store the questions 
# text_elements = pdf.pq('LTTextLineHorizontal') # locate all text elements in the XML file 
# text_elements = pdf.pq('LTTextBoxHorizontal')
# text_elements = pdf.pq('LTTextLineHorizontal')

# for t in text_elements: # loop through each text element 
#     text = t.text # get the text of the element 
#     if len(text) > 0 and text[0].isdigit(): # check if the first character is a digit 
#         questions.append(text) # append the text to the questions list print(questions) # print the questions list
#         print(text)
#     # break
# print(len(questions))


# # 导入PyPDF2模块
# import PyPDF2

# # 打开一个PDF文件对象
# pdf_file = open(pdf_file, 'rb')

# # 创建一个PDF阅读器对象
# pdf_reader = PyPDF2.PdfReader(pdf_file)

# # 获取PDF文件的页数
# num_pages = pdf_reader.pages
# print(num_pages)
# # 遍历每一页
# for page, _ in enumerate(num_pages):
#     # 获取当前页对象
#     pdf_page = pdf_reader.pages[page]
#     # 提取当前页的文本内容
#     page_text = pdf_page.extract_text()
#     # 按换行符分割文本内容为列表
#     lines = page_text.split('\n')
#     # 遍历每一行
#     for line in lines:
#         # 如果行以'ANSWER:'开头，说明是答案
#         if line.startswith('ANSWER:'):
#             # 打印题目和答案
#             print(question)
#             print(line)
#         # 否则，说明是题目或其他内容
#         else:
#             # 保存当前行为题目变量
#             question = line

# # 关闭PDF文件对象
# pdf_file.close()


# 导入python-docx模块
import docx
doc_file = "/home/faith/faith/Downloads/chizong12.docx"
# 打开一个word文档对象
doc = docx.Document(doc_file)
import os
name = os.path.basename(doc_file).replace("docx", '')

print(len(doc.tables))

row_text = []
# 遍历每个表格
for table in doc.tables:
    # 遍历每一行
    for row in table.rows:
        # 创建一个空列表来存储每一行的文本
        
        # 遍历每个单元格
        for cell in row.cells:
            # 提取单元格的文本内容，并添加到列表中
            row_text.append(cell.text)
            next_break = False 
            break_now = False
            for t in cell.tables:
                for r in t.rows:
                    for c in r.cells:
                        if "ANSWER" in c.text:
                            next_break = True
                        row_text.append(c.text)
                        if "ANSWER" not in c.text and next_break:
                            break_now = True
                            break
                    if break_now:
                        break
                if break_now:
                    break
        # 打印每一行的文本内容，用制表符分隔

with open(f'{name}.txt', 'w') as f:
    f.write('\n'.join(row_text)) 
print('\n'.join(row_text))

    
# 遍历每个段落
# for para in doc.paragraphs:
#     # 提取段落的文本内容
#     para_text = para.text
#     print(para_text)
    # # 按换行符分割文本内容为列表
    # lines = para_text.split('\n')
    # # 遍历每一行
    # for line in lines:
    #     # 打印每一行的文本内容
    #     print(line)